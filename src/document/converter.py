"""
문서 변환 모듈
"""

import os
from pathlib import Path
from typing import List, Optional
import logging
from docx import Document
from docx.shared import Inches, Pt

class DocumentConverter:
    """문서 변환 클래스"""
    
    def __init__(self, output_dir: str, title: str = "문서", 
                 image_width: float = 6.0, font_size: int = 11,
                 line_spacing: float = 1.15):
        """
        초기화
        
        Args:
            output_dir: 출력 디렉토리
            title: 문서 제목
            image_width: 이미지 너비 (인치)
            font_size: 폰트 크기
            line_spacing: 줄 간격
        """
        self.output_dir = Path(output_dir)
        self.title = title
        self.image_width = image_width
        self.font_size = font_size
        self.line_spacing = line_spacing
        self.logger = logging.getLogger(__name__)
        
        # 출력 디렉토리 생성
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def create_document(self, images: List[str], texts: List[str],
                       output_file: Optional[str] = None) -> str:
        """
        문서 생성
        
        Args:
            images: 이미지 파일 경로 목록
            texts: 텍스트 목록
            output_file: 출력 파일 경로 (선택)
            
        Returns:
            생성된 문서 파일 경로
        """
        try:
            # 문서 객체 생성
            doc = Document()
            
            # 제목 추가
            doc.add_heading(self.title, 0)
            
            # 이미지와 텍스트 추가
            for image_path, text in zip(images, texts):
                # 이미지 추가
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(self.image_width))
                
                # 텍스트 추가
                paragraph = doc.add_paragraph(text)
                paragraph.style.font.size = Pt(self.font_size)
                paragraph.line_spacing = self.line_spacing
            
            # 출력 파일 경로 설정
            if output_file is None:
                output_file = self.output_dir / f"{self.title}.docx"
            
            # 문서 저장
            doc.save(str(output_file))
            self.logger.info(f"문서 생성 완료: {output_file}")
            
            return str(output_file)
            
        except Exception as e:
            self.logger.error(f"문서 생성 중 오류 발생: {e}")
            raise 