"""
Document conversion module for creating Word documents from images and text.
"""

import os
import logging
from typing import List, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm

class DocumentConverter:
    """문서 변환 클래스"""
    
    def __init__(self, output_path: str, title: str = "Extracted Content"):
        """
        초기화 함수
        
        Args:
            output_path (str): 출력 파일 경로
            title (str): 문서 제목
        """
        self.output_path = Path(output_path)
        self.title = title
        
        # 출력 디렉토리 생성
        os.makedirs(self.output_path.parent, exist_ok=True)
        
        # 로거 설정
        self.logger = logging.getLogger(__name__)
    
    def _add_title(self, doc: Document) -> None:
        """문서 제목 추가"""
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(self.title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        doc.add_paragraph()  # 빈 줄 추가
    
    def _add_image(self, doc: Document, image_path: str) -> None:
        """이미지 추가"""
        try:
            # 이미지 크기 조정
            doc.add_picture(str(image_path), width=Inches(6))
            doc.add_paragraph()  # 빈 줄 추가
        except Exception as e:
            self.logger.error(f"이미지 추가 중 오류 발생: {e}")
    
    def _add_text(self, doc: Document, text: str) -> None:
        """텍스트 추가"""
        if text.strip():
            paragraph = doc.add_paragraph()
            paragraph.add_run(text)
            doc.add_paragraph()  # 빈 줄 추가
    
    def create_document(
        self,
        images: List[str],
        texts: List[str],
        filename: str = "output.docx"
    ) -> Tuple[bool, str]:
        """
        Word 문서 생성
        
        Args:
            images (List[str]): 이미지 파일 경로 목록
            texts (List[str]): 텍스트 목록
            filename (str): 출력 파일명
            
        Returns:
            Tuple[bool, str]: (성공 여부, 출력 파일 경로)
        """
        try:
            # 새 문서 생성
            doc = Document()
            
            # 제목 추가
            self._add_title(doc)
            
            # 이미지와 텍스트 추가
            with tqdm(total=len(images) + len(texts), desc="문서 생성") as pbar:
                # 이미지 추가
                for image_path in images:
                    self._add_image(doc, image_path)
                    pbar.update(1)
                
                # 텍스트 추가
                for text in texts:
                    self._add_text(doc, text)
                    pbar.update(1)
            
            # 문서 저장
            output_file = self.output_path / filename
            doc.save(str(output_file))
            
            self.logger.info(f"문서가 성공적으로 생성되었습니다: {output_file}")
            return True, str(output_file)
            
        except Exception as e:
            self.logger.error(f"문서 생성 중 오류 발생: {e}")
            return False, "" 