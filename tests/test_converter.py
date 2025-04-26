import os
import pytest
import numpy as np
import cv2
from pathlib import Path
from docx import Document
from src.processing.converter import DocumentConverter

class TestDocumentConverter:
    """DocumentConverter 클래스 테스트"""
    
    @pytest.fixture
    def converter(self, tmp_path):
        """DocumentConverter 인스턴스 생성"""
        output_dir = tmp_path / "output"
        return DocumentConverter(output_path=str(output_dir))
    
    @pytest.fixture
    def sample_images(self, tmp_path):
        """테스트용 샘플 이미지 생성"""
        images = []
        
        # 이미지 1 생성
        img1 = np.zeros((100, 100, 3), dtype=np.uint8)
        img1[25:75, 25:75] = 255
        path1 = tmp_path / "image1.png"
        cv2.imwrite(str(path1), img1)
        images.append(str(path1))
        
        # 이미지 2 생성
        img2 = np.zeros((100, 100, 3), dtype=np.uint8)
        img2[50:100, 0:50] = 255
        path2 = tmp_path / "image2.png"
        cv2.imwrite(str(path2), img2)
        images.append(str(path2))
        
        return images
    
    @pytest.fixture
    def sample_texts(self):
        """테스트용 샘플 텍스트"""
        return [
            "첫 번째 텍스트입니다.",
            "두 번째 텍스트입니다."
        ]
    
    def test_create_document(self, converter, sample_images, sample_texts):
        """문서 생성 테스트"""
        # 문서 생성
        success = converter.create_document(
            title="테스트 문서",
            images=sample_images,
            texts=sample_texts,
            output_file="test.docx"
        )
        
        # 문서 생성 성공 확인
        assert success
        
        # 생성된 문서 확인
        doc_path = Path(converter.output_path) / "test.docx"
        assert doc_path.exists()
        
        # 문서 내용 확인
        doc = Document(doc_path)
        assert len(doc.paragraphs) > 0
        assert len(doc.inline_shapes) == len(sample_images)
    
    def test_create_summary_document(self, converter):
        """요약 문서 생성 테스트"""
        # 요약 문서 생성
        success = converter.create_summary_document(
            title="테스트 요약",
            summary="이것은 테스트 요약입니다.",
            key_points=[
                "첫 번째 핵심 포인트",
                "두 번째 핵심 포인트",
                "세 번째 핵심 포인트"
            ],
            output_file="summary.docx"
        )
        
        # 문서 생성 성공 확인
        assert success
        
        # 생성된 문서 확인
        doc_path = Path(converter.output_path) / "summary.docx"
        assert doc_path.exists()
        
        # 문서 내용 확인
        doc = Document(doc_path)
        assert len(doc.paragraphs) > 0
        assert "테스트 요약" in doc.paragraphs[0].text
        assert "요약" in doc.paragraphs[2].text
        assert "핵심 포인트" in doc.paragraphs[4].text
    
    def test_add_image(self, converter):
        """이미지 추가 테스트"""
        # 임시 문서 생성
        doc = Document()
        
        # 이미지 추가
        converter._add_image(
            doc=doc,
            image_path=sample_images[0],
            caption="테스트 이미지"
        )
        
        # 이미지가 추가되었는지 확인
        assert len(doc.inline_shapes) == 1
        assert len(doc.paragraphs) == 2  # 이미지 + 캡션
    
    def test_add_text(self, converter):
        """텍스트 추가 테스트"""
        # 임시 문서 생성
        doc = Document()
        
        # 텍스트 추가
        converter._add_text(doc=doc, text="테스트 텍스트")
        
        # 텍스트가 추가되었는지 확인
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "테스트 텍스트"
    
    def test_add_heading(self, converter):
        """제목 추가 테스트"""
        # 임시 문서 생성
        doc = Document()
        
        # 제목 추가
        converter._add_heading(doc=doc, text="테스트 제목", level=1)
        
        # 제목이 추가되었는지 확인
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "테스트 제목"
    
    def test_error_handling(self, converter):
        """오류 처리 테스트"""
        # 존재하지 않는 이미지 파일
        success = converter.create_document(
            title="테스트",
            images=["nonexistent.jpg"],
            texts=["테스트 텍스트"],
            output_file="error.docx"
        )
        assert not success
        
        # 빈 텍스트 리스트
        success = converter.create_document(
            title="테스트",
            images=[],
            texts=[],
            output_file="empty.docx"
        )
        assert not success 